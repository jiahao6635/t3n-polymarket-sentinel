#!/usr/bin/env python3
"""Build the Google Docs-targeted Terminal 3 submission report."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLACK = RGBColor(0x00, 0x00, 0x00)
MUTED = RGBColor(0x55, 0x55, 0x55)
LINK_BLUE = "1155CC"
FONT = "Arial"


def set_run_font(run, name: str = FONT, size: float | None = None, color=BLACK) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = color


def configure_style(style, *, size: float, color, before: float, after: float, line: float) -> None:
    style.font.name = FONT
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = False
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = style.name.startswith("Heading")
    fmt.widow_control = True


def add_hyperlink(paragraph, text: str, url: str):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    run_properties.extend([fonts, color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([run_properties, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def create_numbering(document: Document, kind: str) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "●" if kind == "bullet" else "%1.")
    level_justification = OxmlElement("w:lvlJc")
    level_justification.set(qn("w:val"), "left")

    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    paragraph_properties.extend([tabs, indent])

    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    run_properties.append(fonts)

    level.extend(
        [start, number_format, level_text, level_justification, paragraph_properties, run_properties]
    )
    abstract.append(level)
    numbering.append(abstract)

    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    num_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number_id = OxmlElement("w:numId")
    number_id.set(qn("w:val"), str(num_id))
    num_properties.extend([level, number_id])
    paragraph_properties.append(num_properties)


def add_list_item(document: Document, num_id: int, text: str, *, bold_prefix: str | None = None):
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    apply_numbering(paragraph, num_id)
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_font(prefix)
        prefix.bold = True
        body = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(body)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_labeled_paragraph(document: Document, label: str, value: str):
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_after = Pt(4)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run)
    label_run.bold = True
    value_run = paragraph.add_run(value)
    set_run_font(value_run)
    return paragraph


def add_linked_labeled_paragraph(document: Document, label: str, text: str, url: str):
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_after = Pt(4)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run)
    label_run.bold = True
    add_hyperlink(paragraph, text, url)
    return paragraph


def add_body(document: Document, text: str):
    paragraph = document.add_paragraph(style="Normal")
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_code_block(document: Document, text: str):
    for line in text.splitlines():
        paragraph = document.add_paragraph(style="Code Evidence")
        run = paragraph.add_run(line or " ")
        set_run_font(run, name="Courier New", size=9, color=BLACK)


def add_finding(document: Document, title: str, evidence: str, recommendation: str):
    document.add_heading(title, level=2)
    paragraph = document.add_paragraph(style="Normal")
    evidence_label = paragraph.add_run("Evidence. ")
    set_run_font(evidence_label)
    evidence_label.bold = True
    evidence_run = paragraph.add_run(evidence)
    set_run_font(evidence_run)
    paragraph = document.add_paragraph(style="Normal")
    recommendation_label = paragraph.add_run("Recommendation. ")
    set_run_font(recommendation_label)
    recommendation_label.bold = True
    recommendation_run = paragraph.add_run(recommendation)
    set_run_font(recommendation_run)


def build(output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_style(
        document.styles["Normal"], size=11, color=BLACK, before=0, after=8, line=1.15
    )
    configure_style(
        document.styles["Heading 1"], size=20, color=BLACK, before=20, after=6, line=1.15
    )
    configure_style(
        document.styles["Heading 2"], size=16, color=BLACK, before=18, after=6, line=1.15
    )
    configure_style(
        document.styles["Heading 3"], size=14, color=MUTED, before=16, after=4, line=1.15
    )

    subtitle = document.styles.add_style("Submission Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    configure_style(subtitle, size=12, color=MUTED, before=0, after=12, line=1.15)
    metadata = document.styles.add_style("Submission Metadata", WD_STYLE_TYPE.PARAGRAPH)
    configure_style(metadata, size=10.5, color=MUTED, before=0, after=3, line=1.15)
    code_style = document.styles.add_style("Code Evidence", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Courier New"
    code_style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Courier New")
    code_style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Courier New")
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = BLACK
    code_style.paragraph_format.left_indent = Inches(0.25)
    code_style.paragraph_format.space_before = Pt(0)
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing = 1.0
    code_style.paragraph_format.widow_control = False

    bullet_id = create_numbering(document, "bullet")

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    title.paragraph_format.line_spacing = 1.0
    title_run = title.add_run("Terminal 3 ADK Onboarding & Contract Deployment")
    set_run_font(title_run, size=26, color=BLACK)
    title_run.bold = False

    subtitle_paragraph = document.add_paragraph(style="Submission Subtitle")
    subtitle_run = subtitle_paragraph.add_run("Polymarket Market-Quality Sentinel | Submission Report")
    set_run_font(subtitle_run, size=12, color=MUTED)

    metadata_paragraph = document.add_paragraph(style="Submission Metadata")
    metadata_run = metadata_paragraph.add_run(
        "Participant: Jiahao (GitHub: jiahao6635) | Verified: 11 August 2026"
    )
    set_run_font(metadata_run, size=10.5, color=MUTED)

    lead = document.add_paragraph(style="Normal")
    lead_label = lead.add_run("Outcome. ")
    set_run_font(lead_label)
    lead_label.bold = True
    lead_text = lead.add_run(
        "Terminal 3 sandbox onboarding, DID authentication, TenantClient validation, custom "
        "Rust/WASM registration, and a live read-only contract invocation all completed successfully."
    )
    set_run_font(lead_text)

    document.add_heading("Submission summary", level=1)
    add_labeled_paragraph(
        document, "DID", "did:t3n:6c90567a5d037e13ae0817b22e6a6fec6630a901"
    )
    add_labeled_paragraph(
        document,
        "Contract",
        "z:6c90567a5d037e13ae0817b22e6a6fec6630a901:pm-sentinel",
    )
    add_labeled_paragraph(document, "Contract ID", "590")
    add_labeled_paragraph(document, "Version", "0.1.0")
    add_labeled_paragraph(document, "Network", "Terminal 3 sandbox")
    add_labeled_paragraph(document, "Registered", "2026-08-11T07:40:42.192Z")
    add_linked_labeled_paragraph(
        document,
        "Repository",
        "github.com/jiahao6635/t3n-polymarket-sentinel",
        "https://github.com/jiahao6635/t3n-polymarket-sentinel",
    )

    document.add_heading("What I built", level=1)
    add_body(
        document,
        "I created a custom Rust/WASM TEE contract instead of stopping at the reference project. "
        "The contract is a read-only Polymarket market-quality sentinel: it accepts normalized public "
        "market data and returns deterministic liquidity, spread, price-consistency, and activity warnings.",
    )
    add_body(
        document,
        "The WASM component imports no host capabilities. It cannot access wallets, secrets, storage, "
        "external networks, or trading APIs. Public Gamma API data is fetched by the TypeScript adapter "
        "outside the contract, which keeps the deployment easy to audit and safe to demonstrate.",
    )

    document.add_heading("Completion status", level=1)
    for item in [
        "Completed: Google SSO claim-page onboarding and one-time API-key storage outside Git.",
        "Completed: DID generation and Terminal 3 sandbox authentication.",
        "Completed: TenantClient validation through tenant.tenant.me().",
        "Completed: Custom Rust contract, six native tests, and wasm32-wasip2 release build.",
        "Completed: Registration of pm-sentinel version 0.1.0 as contract ID 590.",
        "Completed: Live invocation using an active Bitcoin Polymarket market.",
        "Completed: Public GitHub Actions validation for Rust and TypeScript.",
        "Completed: Upstream Terminal 3 bug report and tested pull request.",
    ]:
        add_list_item(document, bullet_id, item, bold_prefix="Completed:")

    document.add_heading("Verified deployment evidence", level=1)
    document.add_heading("Identity and SDK authentication", level=2)
    add_body(
        document,
        "The claim flow generated DID did:t3n:6c90567a5d037e13ae0817b22e6a6fec6630a901. "
        "The adapter authenticated a T3nClient with that DID, constructed a TenantClient with the active "
        "sandbox node URL, and completed tenant.tenant.me() before continuing.",
    )

    document.add_heading("Build and continuous integration", level=2)
    add_body(
        document,
        "Local verification passed cargo fmt, six Rust unit tests, Clippy with warnings denied, the "
        "wasm32-wasip2 release build, and TypeScript type-checking. The same checks passed publicly in "
        "GitHub Actions.",
    )
    add_linked_labeled_paragraph(
        document,
        "CI evidence",
        "GitHub Actions run 31470761322",
        "https://github.com/jiahao6635/t3n-polymarket-sentinel/actions/runs/31470761322",
    )

    document.add_heading("Contract registration", level=2)
    add_body(
        document,
        "Terminal 3 registered script z:6c90567a5d037e13ae0817b22e6a6fec6630a901:pm-sentinel, "
        "version 0.1.0, as numeric contract ID 590 at 2026-08-11T07:40:42.192Z.",
    )

    document.add_heading("Live contract invocation", level=2)
    add_body(
        document,
        "The invocation used market slug will-bitcoin-reach-100k-in-august-2026, an active and "
        "order-accepting market titled 'Will Bitcoin reach $100,000 in August?' with event end date "
        "2026-09-01T04:00:00Z. The contract returned a quality score of 100, deep liquidity, and no warnings.",
    )
    add_code_block(
        document,
        """{
  "question": "Will Bitcoin reach $100,000 in August?",
  "normalized_prices": [0.0015, 0.9985],
  "dominant_outcome": "No",
  "dominant_probability": 0.9985,
  "liquidity_tier": "deep",
  "quality_score": 100,
  "metrics": {
    "price_sum": 1,
    "sum_deviation_bps": 0,
    "spread_bps": 10,
    "liquidity_usd": 140379.14,
    "volume_24h_usd": 33855.55,
    "activity_ratio_bps": 2412
  },
  "warnings": [],
  "disclaimer": "Data-quality screening only; not trading advice."
}""",
    )

    document.add_heading("Engineering findings", level=1)
    add_finding(
        document,
        "A. Quickstart does not type-check against SDK 4.35.0",
        "T3nClientConfig now requires trustAnchor, but the public Quickstart constructor omits it.",
        "Add the required trust-anchor field and explain that unsafe server trust is only a temporary, "
        "non-production escape hatch.",
    )
    add_finding(
        document,
        "B. Development page calls a removed TenantClient method",
        "SDK 4.35.0 exposes the operation as tenant.tenant.me(), not tenant.me(). The documented call "
        "fails during type-checking.",
        "Update the example and expected output to use the namespaced method.",
    )
    add_finding(
        document,
        "C. Reference flight repository documentation is stale",
        "The package is version 0.4.1, while the README previously introduced 0.3.0 and described "
        "inline passenger PII even though the current contract accepts an opaque passenger_id and uses "
        "host-resolved profile placeholders.",
        "Align the README, input example, host capabilities, privacy guarantee, and architecture diagram "
        "with the current source.",
    )
    add_linked_labeled_paragraph(
        document,
        "Public issue",
        "Terminal-3/z-tenant-flight issue 8",
        "https://github.com/Terminal-3/z-tenant-flight/issues/8",
    )
    add_finding(
        document,
        "D. Claim page and Quickstart use different environment names",
        "The claim page demonstrates sandbox, while the Quickstart uses testnet. The onboarding path does "
        "not explain whether these labels share a cluster or have separate DIDs and credits.",
        "Use one label consistently or document the environment mapping.",
    )
    add_finding(
        document,
        "E. Fresh SDK install reports a critical transitive advisory",
        "npm audit identifies public archive-traversal advisories through decompress 4.2.1 in the "
        "componentization toolchain. No exploit attempt was made.",
        "Review whether build-time componentization packages must ship in the runtime dependency graph and "
        "upgrade or constrain the affected dependency where compatible.",
    )
    add_finding(
        document,
        "F. Signed trust-manifest requests return HTTP 405",
        "On 11 August 2026, fetchTrustedManifest for both sandbox and testnet called the official Singapore "
        "node endpoint and failed with 405 Method Not Allowed before handshake.",
        "Fix or publish the signed manifest endpoint before requiring trustAnchor, and update the "
        "Quickstart with a fail-closed example.",
    )
    add_code_block(
        document,
        "Trust manifest request to https://cn-api.sg.testnet.t3n.terminal3.io/api/trust-manifest\n"
        "failed: 405 Method Not Allowed",
    )

    document.add_heading("Upstream contribution", level=1)
    add_body(
        document,
        "I submitted a compatibility fix that normalizes both the public-documentation DID text shape "
        "and the bundled WIT raw 20-byte CompactDid shape before building the tenant secrets-map path. "
        "The patch reuses one helper in both flight operations and adds native regression coverage.",
    )
    add_linked_labeled_paragraph(
        document,
        "Pull request",
        "Terminal-3/z-tenant-flight PR 9",
        "https://github.com/Terminal-3/z-tenant-flight/pull/9",
    )
    for item in [
        "11 native unit tests passed.",
        "Clippy passed with warnings denied.",
        "The wasm32-wasip2 release build passed.",
        "The pull request is open and awaiting maintainer review.",
    ]:
        add_list_item(document, bullet_id, item)

    document.add_heading("Safety and limitations", level=1)
    for item in [
        "No mainnet transaction, wallet connection, deposit, order, or wager was made.",
        "The Polymarket integration is read-only and uses only public Gamma API fields.",
        "The score is a deterministic data-quality heuristic, not a price forecast or trading advice.",
        "T3N_UNSAFE_TRUST_SERVER=1 was used only in Terminal 3 sandbox after explicit authorization because "
        "the signed manifest endpoint returned HTTP 405.",
        "The adapter refuses unsafe trust when T3N_ENV=production.",
        "The one-time Terminal 3 API key remains in a local Git-ignored file and is not included here.",
    ]:
        add_list_item(document, bullet_id, item)

    document.add_heading("Final participant actions", level=1)
    for item in [
        "Leave the one-time API-key success view before capturing any claim-page evidence.",
        "Optionally attach a redacted screenshot showing claim success and DID, with the API key fully hidden.",
        "Review the final bounty text and submit it personally because the listing is marked HUMAN_ONLY.",
        "Use the DID and repository links above in the submission form.",
    ]:
        add_list_item(document, bullet_id, item)

    document.add_heading("Sources", level=1)
    sources = [
        (
            "Project repository",
            "Open the GitHub repository",
            "https://github.com/jiahao6635/t3n-polymarket-sentinel",
        ),
        (
            "Verified CI run",
            "Open GitHub Actions run 31470761322",
            "https://github.com/jiahao6635/t3n-polymarket-sentinel/actions/runs/31470761322",
        ),
        (
            "Terminal 3 issue 8",
            "Open the public bug report",
            "https://github.com/Terminal-3/z-tenant-flight/issues/8",
        ),
        (
            "Terminal 3 pull request 9",
            "Open the upstream pull request",
            "https://github.com/Terminal-3/z-tenant-flight/pull/9",
        ),
        (
            "Polymarket Gamma market record",
            "Open the public market record",
            "https://gamma-api.polymarket.com/markets?slug=will-bitcoin-reach-100k-in-august-2026",
        ),
        (
            "Terminal 3 common integration errors",
            "Open the integration guide",
            "https://docs.terminal3.io/developers/adk/tips/common-errors",
        ),
    ]
    for label, link_text, url in sources:
        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.space_after = Pt(4)
        apply_numbering(paragraph, bullet_id)
        link_label = paragraph.add_run(f"{label}: ")
        set_run_font(link_label)
        link_label.bold = True
        add_hyperlink(paragraph, link_text, url)

    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    document.core_properties.title = "Terminal 3 ADK Onboarding and Contract Deployment"
    document.core_properties.subject = "Polymarket Market-Quality Sentinel submission report"
    document.core_properties.author = "Jiahao"
    document.core_properties.keywords = "Terminal 3, ADK, Rust, WASM, Polymarket"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.output)


if __name__ == "__main__":
    main()
